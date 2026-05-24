from __future__ import annotations

import argparse
import json
import re
import struct
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


FREESECT = 0xFFFFFFFF
ENDOFCHAIN = 0xFFFFFFFE

SPECIAL_8BIT = {
    0x82: 0x201A,
    0x83: 0x0192,
    0x84: 0x201E,
    0x85: 0x2026,
    0x86: 0x2020,
    0x87: 0x2021,
    0x88: 0x02C6,
    0x89: 0x2030,
    0x8A: 0x0160,
    0x8B: 0x2039,
    0x8C: 0x0152,
    0x91: 0x2018,
    0x92: 0x2019,
    0x93: 0x201C,
    0x94: 0x201D,
    0x95: 0x2022,
    0x96: 0x2013,
    0x97: 0x2014,
    0x98: 0x02DC,
    0x99: 0x2122,
    0x9A: 0x0161,
    0x9B: 0x203A,
    0x9C: 0x0153,
    0x9F: 0x0178,
}

DATE_LINE_RE = re.compile(r"^\d{4}年\d{2}月\d{2}日(?:\s+\d{2}:\d{2})?$")
SPEAKER_LINE_RE = re.compile(r"^(?:发言人|说话人|Speaker)\s+\d{2}:\d{2}$")
SHORT_HEADING_RE = re.compile(r"^[^。！？!?]{1,30}[：:]$")
LEADING_INDEX_RE = re.compile(r"^\d+[、，,.\-_ ]*")
DATE_PATTERNS = [
    re.compile(r"[\s_+-]*20\d{6}$"),
    re.compile(r"[\s_+-]*20\d{2}-\d{2}-\d{2}(?:[ _-]?\d{6})?$"),
]


@dataclass
class TranscriptBlock:
    anchor: str
    paragraphs: list[str]


@dataclass
class OutlineSection:
    heading: str
    paragraphs: list[str]


def yaml_quote(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def normalize_text(text: str) -> str:
    text = "".join(ch for ch in text if ch == "\t" or ord(ch) >= 32)
    text = text.replace("\xa0", " ").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def sanitize_filename(name: str) -> str:
    forbidden = '<>:"/\\|?*'
    return "".join("_" if ch in forbidden else ch for ch in name).strip().rstrip(".")


def derive_role_name(file_name: str) -> str:
    stem = Path(file_name).stem
    stem = LEADING_INDEX_RE.sub("", stem)
    previous = None
    while stem != previous:
        previous = stem
        stem = re.sub(r"\s+\(\d+\)$", "", stem).strip()
        stem = stem.replace("_原文", "").replace("原文", "").strip(" _-")
        stem = re.sub(r"\+\d{4}-\d{2}-\d{2}\.mp4(?:_\d{8}_\d{6})?$", "", stem).strip(" _-")
        for pattern in DATE_PATTERNS:
            stem = pattern.sub("", stem).strip(" _-")
        if stem.startswith("《") and stem.endswith("》"):
            stem = stem[1:-1].strip()
        stem = re.sub(r"^任职资格点评[-—_ ]*", "", stem).strip(" _-")
    return stem or Path(file_name).stem


def title_from_role(role_name: str) -> str:
    return f"2026-任职资格点评-{role_name}"


def clean_lines(lines: Iterable[str]) -> list[str]:
    cleaned: list[str] = []
    for line in lines:
        line = normalize_text(line)
        if line:
            cleaned.append(line)
    return cleaned


def iter_docx_blocks(document: DocxDocumentType) -> Iterable[Paragraph | Table]:
    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def table_to_markdown(table: Table) -> list[str]:
    rows = []
    for row in table.rows:
        cells = [normalize_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return []

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def extract_docx_lines(path: Path) -> list[str]:
    document = DocxDocument(path)
    lines: list[str] = []
    table_index = 0

    for block in iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            if text:
                lines.append(text)
        else:
            table_lines = table_to_markdown(block)
            if table_lines:
                table_index += 1
                lines.append(f"表格{table_index}：")
                lines.extend(table_lines)
    return lines


def detect_word_content_type(path: Path) -> str:
    header = path.read_bytes()[:8]
    if header.startswith(b"PK\x03\x04"):
        return "docx"
    if header.startswith(bytes.fromhex("D0CF11E0A1B11AE1")):
        return "doc"
    if header.startswith(b"{\\rtf"):
        return "rtf"
    return path.suffix.lower().lstrip(".")


def extract_rtf_lines(path: Path) -> list[str]:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    raw = raw.replace("{", " ").replace("}", " ")
    return clean_lines(raw.splitlines())


class CompoundFile:
    def __init__(self, data: bytes):
        if data[:8] != bytes.fromhex("D0CF11E0A1B11AE1"):
            raise ValueError("不是有效的 .doc 复合文档格式")
        self.data = data
        self.sector_size = 1 << struct.unpack_from("<H", data, 30)[0]
        self.mini_sector_size = 1 << struct.unpack_from("<H", data, 32)[0]
        self.dir_start = struct.unpack_from("<I", data, 48)[0]
        self.mini_cutoff = struct.unpack_from("<I", data, 56)[0]
        self.minifat_start = struct.unpack_from("<I", data, 60)[0]
        self.csect_minifat = struct.unpack_from("<I", data, 64)[0]
        self.difat_start = struct.unpack_from("<I", data, 68)[0]
        self.csect_difat = struct.unpack_from("<I", data, 72)[0]
        self.fat = self._build_fat()
        self.dir_entries = self._read_directory()
        self.root_entry = next(entry for entry in self.dir_entries if entry["type"] == 5)
        self.mini_stream = self._read_entry_stream(self.root_entry, use_mini=False)
        self.minifat = self._build_minifat()

    def _sector_offset(self, sector_id: int) -> int:
        return (sector_id + 1) * self.sector_size

    def _read_sector(self, sector_id: int) -> bytes:
        offset = self._sector_offset(sector_id)
        return self.data[offset : offset + self.sector_size]

    def _iter_chain(self, start_sector: int, fat: list[int] | None = None) -> Iterable[int]:
        fat = fat or self.fat
        seen: set[int] = set()
        sector_id = start_sector
        while sector_id not in (FREESECT, ENDOFCHAIN) and sector_id < len(fat):
            if sector_id in seen:
                break
            seen.add(sector_id)
            yield sector_id
            sector_id = fat[sector_id]

    def _build_fat(self) -> list[int]:
        difat_entries = list(struct.unpack_from("<109I", self.data, 76))
        fat_sectors = [entry for entry in difat_entries if entry not in (FREESECT, ENDOFCHAIN)]

        next_difat = self.difat_start
        for _ in range(self.csect_difat):
            if next_difat in (FREESECT, ENDOFCHAIN):
                break
            sector_data = self._read_sector(next_difat)
            entries = struct.unpack("<" + "I" * (self.sector_size // 4), sector_data)
            fat_sectors.extend(entry for entry in entries[:-1] if entry not in (FREESECT, ENDOFCHAIN))
            next_difat = entries[-1]

        fat: list[int] = []
        for sector_id in fat_sectors:
            sector_data = self._read_sector(sector_id)
            fat.extend(struct.unpack("<" + "I" * (self.sector_size // 4), sector_data))
        return fat

    def _read_directory(self) -> list[dict]:
        directory_stream = b"".join(self._read_sector(sector) for sector in self._iter_chain(self.dir_start))
        entries = []
        for offset in range(0, len(directory_stream), 128):
            chunk = directory_stream[offset : offset + 128]
            if len(chunk) < 128:
                break
            name_length = struct.unpack_from("<H", chunk, 64)[0]
            name = chunk[: max(0, name_length - 2)].decode("utf-16le", errors="ignore") if name_length >= 2 else ""
            entries.append(
                {
                    "name": name,
                    "type": chunk[66],
                    "start": struct.unpack_from("<I", chunk, 116)[0],
                    "size": struct.unpack_from("<Q", chunk, 120)[0],
                }
            )
        return entries

    def _build_minifat(self) -> list[int]:
        if self.minifat_start in (FREESECT, ENDOFCHAIN) or self.csect_minifat == 0:
            return []
        raw = b"".join(self._read_sector(sector) for sector in self._iter_chain(self.minifat_start))
        return list(struct.unpack("<" + "I" * (len(raw) // 4), raw))

    def _read_entry_stream(self, entry: dict, use_mini: bool = True) -> bytes:
        if entry["size"] == 0:
            return b""

        if use_mini and entry["type"] == 2 and entry["size"] < self.mini_cutoff and self.minifat:
            chunks = []
            sector_id = entry["start"]
            seen: set[int] = set()
            while sector_id not in (FREESECT, ENDOFCHAIN) and sector_id < len(self.minifat):
                if sector_id in seen:
                    break
                seen.add(sector_id)
                offset = sector_id * self.mini_sector_size
                chunks.append(self.mini_stream[offset : offset + self.mini_sector_size])
                sector_id = self.minifat[sector_id]
            return b"".join(chunks)[: entry["size"]]

        chunks = [self._read_sector(sector) for sector in self._iter_chain(entry["start"])]
        return b"".join(chunks)[: entry["size"]]

    def read_stream(self, name: str) -> bytes:
        entry = next((item for item in self.dir_entries if item["name"] == name), None)
        if not entry:
            raise KeyError(name)
        return self._read_entry_stream(entry)


def decode_8bit_text(raw: bytes) -> str:
    return "".join(chr(SPECIAL_8BIT.get(byte, byte)) for byte in raw)


def extract_doc_lines(path: Path) -> list[str]:
    cfb = CompoundFile(path.read_bytes())
    word_stream = cfb.read_stream("WordDocument")

    flags = struct.unpack_from("<H", word_stream, 10)[0]
    table_stream_name = "1Table" if (flags & 0x0200) else "0Table"
    table_stream = cfb.read_stream(table_stream_name)

    csw = struct.unpack_from("<H", word_stream, 32)[0]
    offset = 34 + csw * 2
    cslw = struct.unpack_from("<H", word_stream, offset)[0]
    fib_rg_lw = word_stream[offset + 2 : offset + 2 + cslw * 4]
    ccp_text = struct.unpack_from("<i", fib_rg_lw, 12)[0]

    offset = offset + 2 + cslw * 4
    cb_rg_fc_lcb = struct.unpack_from("<H", word_stream, offset)[0]
    offset += 2
    rg_fc_lcb = word_stream[offset : offset + cb_rg_fc_lcb * 8]
    fc_clx = struct.unpack_from("<I", rg_fc_lcb, 33 * 8)[0]
    lcb_clx = struct.unpack_from("<I", rg_fc_lcb, 33 * 8 + 4)[0]

    clx = table_stream[fc_clx : fc_clx + lcb_clx]
    position = 0
    while position < len(clx) and clx[position] == 0x01:
        cb_grpprl = struct.unpack_from("<h", clx, position + 1)[0]
        position += 3 + cb_grpprl

    if position >= len(clx) or clx[position] != 0x02:
        raise ValueError(f"{path.name} 的 CLX 结构中未找到 Piece Table")

    lcb_piece_table = struct.unpack_from("<I", clx, position + 1)[0]
    plc_pcd = clx[position + 5 : position + 5 + lcb_piece_table]

    piece_count = (lcb_piece_table - 4) // 12
    cps = list(struct.unpack_from("<" + "I" * (piece_count + 1), plc_pcd, 0))
    pcd_base = 4 * (piece_count + 1)

    parts: list[str] = []
    for index in range(piece_count):
        cp_start = cps[index]
        cp_end = cps[index + 1]
        if cp_start >= ccp_text:
            break
        char_count = min(cp_end, ccp_text) - cp_start
        if char_count <= 0:
            continue

        pcd = plc_pcd[pcd_base + index * 8 : pcd_base + (index + 1) * 8]
        fc_raw = struct.unpack_from("<I", pcd, 2)[0]
        compressed = (fc_raw >> 30) & 1
        fc = fc_raw & 0x3FFFFFFF

        if compressed:
            byte_start = fc // 2
            chunk = word_stream[byte_start : byte_start + char_count]
            parts.append(decode_8bit_text(chunk))
        else:
            byte_start = fc
            chunk = word_stream[byte_start : byte_start + char_count * 2]
            parts.append(chunk.decode("utf-16le", errors="ignore"))

    text = "".join(parts)
    replacements = {
        "\r": "\n",
        "\x07": "\n",
        "\x0b": "\n",
        "\x0c": "\n",
        "\x13": "",
        "\x14": "",
        "\x15": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    return clean_lines(text.split("\n"))


def looks_like_title(line: str, role_name: str) -> bool:
    if len(line) > 80:
        return False
    if line.endswith(("。", "！", "？", "；", "，", ".", "!", "?", ";", ",")):
        return False
    line_simplified = normalize_text(line).replace("《", "").replace("》", "")
    role_simplified = role_name.replace("《", "").replace("》", "")
    return role_simplified and role_simplified in line_simplified


def is_transcript(lines: list[str]) -> bool:
    return sum(1 for line in lines if SPEAKER_LINE_RE.match(line)) >= 3


def parse_transcript(lines: list[str]) -> tuple[str | None, str | None, list[str], list[TranscriptBlock]]:
    raw_title = None
    doc_date = None
    position = 0

    while position < len(lines):
        line = lines[position]
        if raw_title is None and not SPEAKER_LINE_RE.match(line) and not DATE_LINE_RE.match(line):
            raw_title = line
            position += 1
            continue
        if doc_date is None and DATE_LINE_RE.match(line):
            doc_date = line
            position += 1
            continue
        break

    preface: list[str] = []
    blocks: list[TranscriptBlock] = []
    current_anchor: str | None = None
    current_paragraphs: list[str] = []

    for line in lines[position:]:
        if SPEAKER_LINE_RE.match(line):
            if current_anchor:
                blocks.append(TranscriptBlock(current_anchor, current_paragraphs[:]))
            elif current_paragraphs:
                preface.extend(current_paragraphs)
            current_anchor = line
            current_paragraphs = []
        else:
            current_paragraphs.append(line)

    if current_anchor:
        blocks.append(TranscriptBlock(current_anchor, current_paragraphs[:]))
    elif current_paragraphs:
        preface.extend(current_paragraphs)

    return raw_title, doc_date, preface, blocks


def parse_outline(lines: list[str], role_name: str) -> tuple[str | None, str | None, list[OutlineSection]]:
    raw_title = None
    doc_date = None
    index = 0

    if index < len(lines) and looks_like_title(lines[index], role_name):
        raw_title = lines[index]
        index += 1
    if index < len(lines) and DATE_LINE_RE.match(lines[index]):
        doc_date = lines[index]
        index += 1

    sections: list[OutlineSection] = []
    current_heading = "正文原文"
    current_paragraphs: list[str] = []

    for line in lines[index:]:
        if SHORT_HEADING_RE.match(line):
            if current_paragraphs:
                sections.append(OutlineSection(current_heading.rstrip("：:"), current_paragraphs[:]))
            current_heading = line
            current_paragraphs = []
        else:
            current_paragraphs.append(line)

    if current_paragraphs:
        sections.append(OutlineSection(current_heading.rstrip("：:"), current_paragraphs[:]))

    return raw_title, doc_date, sections


def build_index_pairs_from_transcript(role_name: str, source_file: str, doc_date: str | None, blocks: list[TranscriptBlock]) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = [
        (f"这份任职资格点评文档对应的岗位是什么？", role_name),
        (f"{role_name} 这份文档的原始文件名是什么？", source_file),
    ]
    if doc_date:
        pairs.append((f"{role_name} 这份文档记录的日期是什么？", doc_date))

    for block in blocks[:5]:
        answer = "\n".join(block.paragraphs).strip()
        if answer:
            pairs.append((f"{role_name} 文档在“{block.anchor}”这一段讲了什么？", answer))
    return pairs[:8]


def build_index_pairs_from_outline(role_name: str, source_file: str, doc_date: str | None, sections: list[OutlineSection]) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = [
        ("这份任职资格点评文档对应的岗位是什么？", role_name),
        (f"{role_name} 这份文档的原始文件名是什么？", source_file),
    ]
    if doc_date:
        pairs.append((f"{role_name} 这份文档记录的日期是什么？", doc_date))

    for section in sections[:5]:
        answer = "\n".join(section.paragraphs).strip()
        if answer:
            pairs.append((f"{role_name} 文档中“{section.heading}”部分写了什么？", answer))
    return pairs[:8]


def build_markdown(path: Path, role_name: str, lines: list[str]) -> str:
    title = title_from_role(role_name)
    output_lines = [
        "---",
        f"title: {yaml_quote(title)}",
        f"source_file: {yaml_quote(path.name)}",
        f"source_path: {yaml_quote(str(path))}",
        f"source_format: {yaml_quote(path.suffix.lstrip('.').lower())}",
        'topology: "流式文本"',
        "tags:",
        '  - "任职资格点评"',
        f"  - {yaml_quote(role_name)}",
        '  - "原文转写"',
        "rag_optimized: true",
        "---",
        "",
        f"## 任职资格点评原文：{role_name}",
        "",
    ]

    if is_transcript(lines):
        raw_title, doc_date, preface, blocks = parse_transcript(lines)
        if raw_title:
            output_lines.extend(["### 原始标题", "", raw_title, ""])
        if doc_date:
            output_lines.extend(["### 原始日期", "", doc_date, ""])
        if preface:
            output_lines.extend(["### 前置原文", ""])
            for paragraph in preface:
                output_lines.extend([paragraph, ""])
        for block in blocks:
            output_lines.extend([f"### {block.anchor}", ""])
            for paragraph in block.paragraphs:
                output_lines.extend([paragraph, ""])
        index_pairs = build_index_pairs_from_transcript(role_name, path.name, doc_date, blocks)
    else:
        raw_title, doc_date, sections = parse_outline(lines, role_name)
        if raw_title:
            output_lines.extend(["### 原始标题", "", raw_title, ""])
        if doc_date:
            output_lines.extend(["### 原始日期", "", doc_date, ""])
        for section in sections:
            output_lines.extend([f"### {section.heading}", ""])
            for paragraph in section.paragraphs:
                output_lines.extend([paragraph, ""])
        index_pairs = build_index_pairs_from_outline(role_name, path.name, doc_date, sections)

    output_lines.extend([f"## RAG智能索引：{role_name}", ""])
    for question, answer in index_pairs:
        output_lines.extend([f"Q: {question}", f"A: {answer}", ""])

    return "\n".join(output_lines).rstrip() + "\n"


def convert_file(path: Path, output_dir: Path) -> tuple[Path, str]:
    role_name = derive_role_name(path.name)
    content_type = detect_word_content_type(path)
    if content_type == "docx":
        lines = extract_docx_lines(path)
    elif content_type == "doc":
        lines = extract_doc_lines(path)
    elif content_type == "rtf":
        lines = extract_rtf_lines(path)
    else:
        raise ValueError(f"暂不支持的文档内部格式：{content_type}")
    markdown = build_markdown(path, role_name, lines)
    title = title_from_role(role_name)
    output_path = output_dir / f"{sanitize_filename(title)}.md"
    output_path.write_text(markdown, encoding="utf-8")
    return output_path, title


def collect_source_files(input_dir: Path) -> list[Path]:
    return sorted(
        [
            path
            for path in input_dir.iterdir()
            if path.is_file() and path.suffix.lower() in {".doc", ".docx"} and not path.name.startswith("~$")
        ],
        key=lambda path: path.name,
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="批量将 Word 文档转为 RAG 友好的 Markdown。")
    parser.add_argument("input_dir", help="待转换文档所在目录")
    parser.add_argument("--output-dir", help="Markdown 输出目录，默认与原文档同目录")
    args = parser.parse_args()

    input_dir = Path(args.input_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else input_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    files = collect_source_files(input_dir)
    if not files:
        print("没有找到可转换的 .doc 或 .docx 文件。")
        return 1

    print(f"找到 {len(files)} 份 Word 文档，开始转换。")
    success_count = 0
    failed: list[tuple[str, str]] = []

    for path in files:
        try:
            output_path, title = convert_file(path, output_dir)
            success_count += 1
            print(f"[成功] {path.name} -> {output_path.name} ({title})")
        except Exception as exc:
            failed.append((path.name, str(exc)))
            print(f"[失败] {path.name}: {exc}")

    print("")
    print(f"转换完成：成功 {success_count} 份，失败 {len(failed)} 份。")
    if failed:
        print("失败清单：")
        for name, reason in failed:
            print(f"- {name}: {reason}")
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
